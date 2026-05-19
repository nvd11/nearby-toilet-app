import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    return h

# Header
h1 = add_heading('Wenlin Pan (Jason)', 0)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Target Position: Model Validation Director / Head of AI Safety & Governance / AI Agent & Cloud Architect\n').bold = True
p.add_run('Location: Guangzhou, China | Experience: 17 Years (10+ Years in FinTech & Risk Compliance)\n')
p.add_run('Education: Bachelor of Computer Science, Guangdong University of Foreign Studies (2003-2007)')

# Summary
add_heading('Professional Summary', 1)
doc.add_paragraph('Senior FinTech & AI Architecture Expert with 17 years of software architecture and development experience, deeply specialized in Financial Risk & Compliance for over 10 years.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Cloud Tech & Architecture: ').bold = True
p.add_run('Deep expertise in Google Cloud Platform (GCP) architecture. Proficient in Dataflow (Apache Beam), BigQuery, Pub/Sub, and Cloud Run to build high-concurrency, high-availability Streaming ETL pipelines and microservice architectures. Proven track record of building enterprise-grade data platforms and MLOps platforms from scratch.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('AI Agent & Model Evaluation: ').bold = True
p.add_run('Focused on the enterprise-level implementation and governance of Generative AI (GenAI) and Large Language Models (LLMs). Led the end-to-end lifecycle management of the compliance AI assistant (Project ADA) based on RAG architecture. Highly proficient in AI Agent Orchestration and deeply experienced in Model Evaluation, having established an automated evaluation framework covering completeness, faithfulness, accuracy, and toxicity detection (e.g., LLM-as-a-Judge) to ensure compliance with Responsible AI standards.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Data Engineering & Governance: ').bold = True
p.add_run('Responsible for the construction and evolution of HSBC\'s global Regulatory Compliance Data Platform (RCDP). Expert in integrating massive structured and unstructured data, defining data dictionaries, implementing data access segregation (BQ Auth Views), and enforcing data sensitivity tiering strategies. Ensures data processing and model applications comply with global financial regulatory standards (e.g., SR 11-7, GDPR).')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Technical Leadership: ').bold = True
p.add_run('Experienced in managing global R&D teams. Adept at leveraging Agile methodologies to bridge Data Scientists, Risk Compliance experts, and Software Engineers to drive the secure and compliant delivery of complex AI and data engineering projects.')

# Skills
add_heading('Core Skills', 1)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cloud Computing & MLOps: ').bold = True
p.add_run('Google Cloud Platform (GCP), Dataflow (Apache Beam), BigQuery, Pub/Sub, Cloud Run, Kubernetes (K8s/GKE), Kubeflow, Terraform (IaC).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('AI & Model Validation: ').bold = True
p.add_run('AI Agent Development (RAG Architecture), LLM Evaluation & Monitoring Frameworks, Prompt Engineering Governance, Vector Databases (AlloyDB).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Data Engineering & Governance: ').bold = True
p.add_run('Batch & Streaming ETL, Data Lifecycle Management, Data Sensitivity Tiering & Entitlement Engines, Regulatory Compliance Data Standards.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Programming & Frameworks: ').bold = True
p.add_run('Python (Expert), Java (Expert), SQL, Kubernetes (K8s)-based Microservices Architecture & Development, LangChain, Spring Boot, FastAPI, Django.')

# Experience
add_heading('Work Experience & Key Projects', 1)

doc.add_heading('HSBC Software Development', 2)
p = doc.add_paragraph()
p.add_run('Position: ').bold = True
p.add_run('Technical Lead / R&D Manager | ')
p.add_run('Date: ').bold = True
p.add_run('2018.06 - Present')
doc.add_paragraph('Business Context: Responsible for the global Regulatory Compliance Data Platform (RCDP). RCDP is the core data hub for the compliance department, utilizing GCP Dataflow to ingest and process massive upstream regulatory data, empowering downstream systems (e.g., CMLP Machine Learning Platform, Compliance Q&A System) with high-quality data and AI capabilities via BigQuery and REST APIs.')

# Project 1
doc.add_heading('Key Project 1: Project ADA (Ask Compliance Digital Assistant) - AI Agent / RAG & Model Evaluation', 3)
doc.add_paragraph('Background: A compliance Q&A digital assistant for global front-line bank users, designed to automate policy answering via AI, drastically reducing manual review costs. The project consists of four core subsystems: Ask Compliance (Front-end), RCDP (Data & Orchestration), GenAI (Model provider), and GPPS (Policy repository).\nRole & Responsibilities: RCDP Engineering & Tech Lead.')

doc.add_paragraph('RAG Architecture & Agent Orchestration:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('Designed and implemented the Orchestration Service based on microservices. Upon receiving streaming query data via Pub/Sub, RCDP leverages Python and BM25 algorithms for intent recognition, sensitivity detection, and relevance scoring, seamlessly integrating with GenAI\'s RAG interface upon passing checks.', style='List Bullet 2')
doc.add_paragraph('Drove the construction of a high-quality vector knowledge base based on GCP AlloyDB, building automated update pipelines for GPPS policy documents to ensure embedding accuracy and timeliness.', style='List Bullet 2')

doc.add_paragraph('Automated Model Evaluation & Monitoring:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('Built a continuous evaluation and monitoring workflow from scratch for LLM responses. Utilized BigQuery to capture original user queries, AI responses, referenced chunk match scores, and final user feedback (closed-loop data collection).', style='List Bullet 2')
doc.add_paragraph('Collaborated with Data Scientists to develop an Evaluation API. Implemented core metrics calculation (Completeness, Faithfulness, Accuracy, Toxicity) incorporating LLM-as-a-Judge mechanisms.', style='List Bullet 2')
doc.add_paragraph('Pushed multi-dimensional evaluation metrics to MI systems for reporting, providing a quantitative basis for prompt optimization and RAG framework iterations, meeting strict internal Model Risk Management (MRM) requirements.', style='List Bullet 2')

# Project 2
doc.add_heading('Key Project 2: Rapid Alert Auto Triage - Streaming ETL & MLOps', 3)
doc.add_paragraph('Background: The Rapid2 system receives massive regulatory alerts daily. This project introduces ML models to automate alert triage, saving human analysis costs.\nRole & Responsibilities: Data Engineering & MLOps Architect.')

doc.add_paragraph('Data Engineering & Strict Data Governance:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('Built robust batch (Daily) and Streaming ETL data pipelines to extract historical and incremental Rapid2 data into RCDP BigQuery.', style='List Bullet 2')
doc.add_paragraph('Implemented bank-grade data access control: based on data dictionaries (Public/Internal/Restricted/Highly Restricted) and internal entitlement engines to dynamically generate secure BQ Auth Views for downstream data scientists, ensuring zero sensitive data leakage during model training.', style='List Bullet 2')

doc.add_paragraph('MLOps Platform & Productionization:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('Led the development of the Model Execution Environment to solve the pain point of data scientists lacking productionization experience.', style='List Bullet 2')
doc.add_paragraph('Built automated CI/CD Pipelines (Jenkins/Ansible) to support one-click deployment of ML models from CMLP (GCP Kubeflow/Jupyter) to the execution environment, achieving automated testing, release, and production monitoring.', style='List Bullet 2')

doc.add_paragraph('Event-Driven Architecture & Decoupling:', style='List Bullet').runs[0].bold = True
doc.add_paragraph('Established bi-directional streaming pipelines: Rapid2 pushes new alerts to RCDP Pub/Sub; RCDP processes via Streaming ETL and invokes the AI model in the execution environment; model recommendations are sent back to the user via a real-time Reverse Flow.', style='List Bullet 2')
doc.add_paragraph('Established an Evaluation Flow to daily compare model recommendations against actual business triage decisions, generating evaluation reports and enabling "zero-code-intrusion" AI capabilities for the business system.', style='List Bullet 2')

# Other Projects
doc.add_heading('Other Key Projects', 3)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Global Conflict Management System (GCMS): ').bold = True
p.add_run('China Team Manager / Pod Lead (Managed 9 members). Led Agile transformation (2-week cycle), introduced CI/CD, feature toggles, and blue-green deployments. Migrated to Spring Cloud microservices deployed on Kubernetes (K8s), reaching 70% automated test coverage. (2021.08 - 2023.04)')

p = doc.add_paragraph(style='List Bullet')
p.add_run('QuickApp (EUC Modernization Platform): ').bold = True
p.add_run('Pod Lead. Designed Python Django Web architecture supporting self-service deployment, successfully refactoring 10+ high-risk Excel macros into compliance systems. (2021.02 - 2021.08)')

# Previous Experience
doc.add_heading('Previous Experience', 2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('AIA Information Technology | Senior Software Engineer | 2017.11 - 2018.06\n').bold = True
p.add_run('Responsible for the migration and validation of the core Group Insurance system, ensuring data consistency and business logic accuracy between legacy and new systems.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Shanghai Wicresoft (Client: AIA) | System Engineer | 2016.12 - 2017.11\n').bold = True
p.add_run('Led technical design and Code Review of Group Insurance modules, provided UAT support, and ensured high system availability.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Tata Consultancy Services (Client: HSBC) | Senior Software Engineer | 2011.03 - 2016.12\n').bold = True
p.add_run('Participated in the R&D of HSBC\'s global core trading systems. Designed high-concurrency transaction processing modules and implemented strict unit/integration testing.')

doc.save('/home/gateman/.openclaw/workspace/temp_resume/Jason_Pan_Resume_EN.docx')
print("English Docx created successfully.")
