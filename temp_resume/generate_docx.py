import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    return h

# Header
h1 = add_heading('潘文林 (Jason)', 0)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('求职意向：Model Validation Director / Head of AI Safety & Governance / AI Agent & Cloud Architect\n').bold = True
p.add_run('地点：中国广州 | 经验：17年 (10+年金融科技与风险合规) | 教育：广东外语外贸大学 - 计算机科学与技术 学士 (2003-2007)')

# Summary
add_heading('个人简介 (Professional Summary)', 1)
doc.add_paragraph('拥有17年软件架构与开发经验的资深金融科技与AI架构专家，在金融风险与合规（Risk & Compliance）领域深耕10年以上。')

p = doc.add_paragraph(style='List Bullet')
p.add_run('GCP 云计算与架构能力 (Cloud Tech & Architecture)：').bold = True
p.add_run('拥有深厚的 Google Cloud Platform (GCP) 公有云架构经验。熟练运用 Dataflow (Apache Beam)、BigQuery、Pub/Sub、Cloud Run 构建高并发、高可用的流式数据处理 (Streaming ETL) 和微服务架构。具备从零构建企业级数据中台及 MLOps 平台的能力。')

p = doc.add_paragraph(style='List Bullet')
p.add_run('AI Agent 开发与模型评估 (AI Agent & Model Evaluation)：').bold = True
p.add_run('专注于生成式 AI (GenAI) 和大型语言模型 (LLM) 的企业级落地与治理。主导基于 RAG (Retrieval-Augmented Generation) 架构的合规 AI 助手 (Project ADA) 的端到端生命周期管理。精通 AI Agent 编排机制 (Orchestration)，并在模型评估 (Model Evaluation) 领域经验丰富，建立了一套涵盖准确性、合规性及毒性检测的自动化评估框架（如 LLM-as-a-Judge），确保系统符合负责任的 AI (Responsible AI) 标准。')

p = doc.add_paragraph(style='List Bullet')
p.add_run('数据工程与数据治理 (Data Engineering & Governance)：').bold = True
p.add_run('负责汇丰银行全球合规数据平台 (RCDP) 的建设与演进。精通海量结构化与非结构化数据的处理与整合，制定并实施了严格的数据字典、数据权限隔离 (BQ Auth Views) 及数据敏感度分级策略。确保数据处理和模型应用符合全球金融监管标准（如 SR 11-7, GDPR）。')

p = doc.add_paragraph(style='List Bullet')
p.add_run('技术领导力 (Technical Leadership)：').bold = True
p.add_run('具备全球化研发团队的管理经验，擅长运用敏捷开发体系，连接数据科学家 (Data Scientists)、风控业务专家及研发工程师，推动复杂 AI 项目和数据工程的安全与合规交付。')

# Skills
add_heading('核心技能 (Core Skills)', 1)
p = doc.add_paragraph(style='List Bullet')
p.add_run('云计算与 MLOps：').bold = True
p.add_run('Google Cloud Platform (GCP), Dataflow (Apache Beam), BigQuery, Pub/Sub, Cloud Run, Kubernetes (K8s/GKE), Kubeflow, Terraform (IaC)。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('AI 与模型验证：').bold = True
p.add_run('AI Agent 开发 (RAG 架构), 大模型评估体系 (LLM Evaluation & Monitoring), 提示词工程治理, 向量数据库 (AlloyDB)。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('数据工程与治理：').bold = True
p.add_run('批处理与流式数据处理 (Streaming ETL), 数据生命周期管理, 数据敏感度分级管控与权限引擎, 监管合规数据标准。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('编程语言与框架：').bold = True
p.add_run('Python (资深), Java (资深), SQL, 基于 Kubernetes (K8s) 的微服务架构与开发, LangChain, Spring Boot, FastAPI, Django。')

# Experience
add_heading('工作履历与核心项目经验 (Work Experience & Key Projects)', 1)

doc.add_heading('汇丰软件开发（广东）有限公司 (HSBC Software Development)', 2)
p = doc.add_paragraph()
p.add_run('职位：').bold = True
p.add_run('Technical Lead / R&D Manager | ')
p.add_run('时间：').bold = True
p.add_run('2018.06 - 至今')
doc.add_paragraph('业务背景：负责汇丰银行全球风险合规数据平台 (Regulatory Compliance Data Platform, RCDP)。RCDP 是合规部门的核心数据中枢，利用 GCP Dataflow 收集和处理海量上游监管数据，并通过 BigQuery 和 REST API 为下游系统（如 CMLP 机器学习平台、合规问答系统）提供高质量的数据支持与 AI 能力赋能。')

# Project 1
doc.add_heading('核心项目 1：Project ADA (Ask Compliance Digital Assistant) - AI Agent / RAG 与模型评估', 3)
doc.add_paragraph('项目背景：面向全球银行前台用户的合规问答数字助理。旨在通过 AI 技术自动解答合规政策问题，大幅降低合规专家的查阅成本。项目由四个核心子系统协作完成：Ask Compliance (前端), RCDP (数据与编排), GenAI (模型端) 以及 GPPS (政策库)。\n角色与职责：RCDP 研发及技术负责人')

doc.add_paragraph('RAG 架构设计与 Agent 编排 (AI Agent Development)：', style='List Bullet').runs[0].bold = True
doc.add_paragraph('设计并实现了基于微服务架构的编排服务 (Orchestration Service)。RCDP 接收流式查询数据 (Pub/Sub) 后，利用 Python 与 BM25 算法进行意图识别、敏感度检测及相关性判断。评估合格后无缝对接 GenAI 的 RAG 接口。', style='List Bullet 2')
doc.add_paragraph('推动基于 GCP AlloyDB 的高质量向量知识库建设。打通 GPPS 政策文档的自动化更新链路，通过 RCDP 不断将最新政策推送到 GenAI 端进行 Embedding 生成，确保底层知识库的准确性和时效性。', style='List Bullet 2')

doc.add_paragraph('自动化模型评估体系建设 (Model Evaluation & Monitoring)：', style='List Bullet').runs[0].bold = True
doc.add_paragraph('从零构建了针对大模型回答的持续评估与监控工作流 (Model Evaluation Flow)。利用 BigQuery 捕获用户原始查询、AI 回答、参考的 chunk 匹配度分数以及用户最终反馈（闭环数据收集）。', style='List Bullet 2')
doc.add_paragraph('与数据科学家合作开发 Evaluation API。实施涵盖 Completeness (完整性), Faithfulness (忠实度), Accuracy (准确率) 及 Toxicity (毒性) 在内的核心指标计算（包含 LLM-as-a-Judge 机制）。', style='List Bullet 2')
doc.add_paragraph('将生成的多维评估指标推送至 MI (管理信息) 系统进行报表展现，为后续提示词优化和 RAG 框架的持续迭代提供量化依据，满足银行内部严苛的模型风险管理 (MRM) 要求。', style='List Bullet 2')

# Project 2
doc.add_heading('核心项目 2：Rapid Alert Auto Triage - 流式处理 ETL 与 MLOps', 3)
doc.add_paragraph('项目背景：汇丰合规部门使用的 Rapid2 系统每天接收海量待分类的监管警报 (Regulatory Alerts)。本项目引入机器学习模型对警报进行自动化分流处理，大幅节省分析成本。\n角色与职责：数据工程与 MLOps 架构师')

doc.add_paragraph('数据工程与严格数据治理 (Data Engineering & Governance)：', style='List Bullet').runs[0].bold = True
doc.add_paragraph('构建稳健的批处理 (Daily Batch) 与流处理 (Streaming ETL) 数据管道。将 Rapid2 历史与增量数据抽取到 RCDP BigQuery 中。', style='List Bullet 2')
doc.add_paragraph('实施银行级数据权限分级控制：基于数据字典（Public / Internal / Restricted / Highly Restricted），结合内部权限引擎，为下游数据科学家动态生成安全的 BigQuery Auth Views，确保模型训练时敏感数据零泄露。', style='List Bullet 2')

doc.add_paragraph('MLOps 平台建设与模型工程化 (MLOps & Productionization)：', style='List Bullet').runs[0].bold = True
doc.add_paragraph('主导开发了 Model Execution Environment (模型执行环境)，解决数据科学家缺乏生产化经验的痛点。', style='List Bullet 2')
doc.add_paragraph('搭建自动化 CI/CD Pipeline (Jenkins/Ansible)：支持将 CMLP (基于 GCP Kubeflow / Jupyter) 上的 ML 模型一键部署至执行环境。实现自动化测试、发布以及生产环境监控。', style='List Bullet 2')

doc.add_paragraph('业务解耦与双向流式集成 (Event-Driven Architecture)：', style='List Bullet').runs[0].bold = True
doc.add_paragraph('建立双向 Streaming 流水线：Rapid2 将新增警报推送至 RCDP Pub/Sub Topic；RCDP 通过流式 ETL 处理后调用执行环境中的 AI 模型；模型推荐 (Recommendation) 再通过另一条下行通道 (Reverse Flow) 实时送回用户端。', style='List Bullet 2')
doc.add_paragraph('建立模型评估回路 (Evaluation Flow)，每日批量比对模型推荐结果与业务人员实际分类决策，生成评估报表，实现对业务系统“无代码侵入”的 AI 能力赋能。', style='List Bullet 2')

# Other Projects
doc.add_heading('其他核心项目', 3)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Global Conflict Management System (GCMS): ').bold = True
p.add_run('中国团队负责人(Pod Lead，管理9人)。带领团队完成敏捷转型，缩短交付周期至2周。引入CI/CD自动化部署、特性开关(Feature Toggles)和蓝绿部署，完成向 Spring Cloud 微服务架构的迁移，基于 Kubernetes (K8s) 部署微服务，自动化测试覆盖率达70%。(2021.08 - 2023.04)')

p = doc.add_paragraph(style='List Bullet')
p.add_run('QuickApp (EUC 现代化迁移平台): ').bold = True
p.add_run('Pod Lead。设计 Python Django Web 架构，支持用户自助部署应用，成功重构 10+ 个高风险 Excel 宏工具为合规系统。(2021.02 - 2021.08)')

# Previous Experience
doc.add_heading('早期职业经历 (Previous Experience)', 2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('AIA Information Technology (友邦保险) | Senior Software Engineer | 2017.11 - 2018.06\n').bold = True
p.add_run('负责核心团险系统的迁移和验证，确保遗留系统与新系统之间的数据一致性和业务逻辑准确性，监控核心业务性能指标。')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Shanghai Wicresoft (外派至 AIA) | System Engineer | 2016.12 - 2017.11\n').bold = True
p.add_run('主导团险业务模块的技术设计与 Code Review，提供 UAT 支持，处理生产环境疑难问题，保障系统高可用。')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Tata Consultancy Services (外派至汇丰) | Senior Software Engineer | 2011.03 - 2016.12\n').bold = True
p.add_run('参与汇丰全球核心交易系统研发，负责高并发事务处理模块设计，实施严格的单元和集成测试，确保海量金融数据处理的绝对准确。')

doc.save('/home/gateman/.openclaw/workspace/temp_resume/Jason_Pan_Resume_V2.docx')
print("Docx V2 created successfully.")
