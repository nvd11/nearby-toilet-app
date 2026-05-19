import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h

# Header
h = add_heading('Jason Pan (Wenlin Pan)', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Target Roles: ').bold = True
p.add_run('Head of Data / Cloud Data Architect / Senior Data Engineer\n')
p.add_run('Location: ').bold = True
p.add_run('Guangzhou, China | ')
p.add_run('Experience: ').bold = True
p.add_run('17 Years (10+ Years in FinTech Data Platforms & GCP Cloud-Native Architecture)\n')
p.add_run('Education: ').bold = True
p.add_run('Bachelor of Computer Science, Guangdong University of Foreign Studies')

add_heading('Executive Summary', level=2)
p = doc.add_paragraph('Senior Data & Cloud Infrastructure Architect with 17+ years of software engineering experience. Expert in building PB-scale enterprise data platforms on GCP, driving overall strategic migrations from legacy on-premise data centers to modern cloud-native architectures.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Enterprise Data Platforms & Cloud Migration: ').bold = True
p.add_run('Led the architecture and engineering of HSBC\'s Regulatory Compliance Data Platform (RCDP). Formulated and executed the "CDR Demise" strategy, successfully migrating core business data from a legacy on-premise Oracle data warehouse to a modern GCP data architecture.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Data Governance & Security: ').bold = True
p.add_run('Deep understanding of financial data security and compliance (GDPR, banking secrecy laws). Designed and implemented an automated Entitlement Engine that dynamically generates BigQuery Auth Views based on field-level sensitivity (Public/Restricted/etc.), ensuring strict fine-grained access control.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Streaming & Batch ETL Pipelines: ').bold = True
p.add_run('Expert in high-throughput data processing using Apache Beam (Dataflow). Built dual-track data pipelines accommodating both low-latency (Pub/Sub + Dataflow) and high-throughput batch processing (Cloud Scheduler + Dataflow + Cloud Storage) requirements.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cloud Infrastructure & MLOps: ').bold = True
p.add_run('Proficient in underlying GCP infrastructure, including VPC network planning, IAM security policies, and GKE containerization. Built a Model Execution Environment from scratch, bridging the gap between data science and engineering by automating the MLOps lifecycle from data ingestion to model deployment and monitoring.')

add_heading('Core Competencies', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Data Engineering: ').bold = True
p.add_run('Dataflow (Apache Beam), Streaming/Batch ETL, Pub/Sub, Cloud Storage, High-concurrency Data Ingestion.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Data Governance: ').bold = True
p.add_run('Data Lineage, Data Sensitivity Classification, BigQuery Auth Views, Entitlement Engine Design, Principle of Least Privilege (PoLP).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cloud Infra (GCP): ').bold = True
p.add_run('BigQuery, GKE (Kubernetes), Cloud Run, VPC/IAM/Service Accounts, Cloud Scheduler, Terraform (IaC).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Software Engineering: ').bold = True
p.add_run('Python (Asynchronous Programming), Java (Spring Boot/Spring Cloud), FastAPI, RESTful API Architecture, CI/CD (Jenkins/Cloud Build).')


add_heading('Professional Experience', level=2)
add_heading('HSBC Software Development (Guangdong)', level=3)
p = doc.add_paragraph()
p.add_run('Role: ').bold = True
p.add_run('Technical Lead & Engineering Manager - Regulatory Compliance Data Platform (RCDP) | ')
p.add_run('Date: ').bold = True
p.add_run('2018.06 - Present\n')
p.add_run('Core Responsibility: ').bold = True
p.add_run('Fully responsible for the technology stack selection, infrastructure architecture, and core framework development of RCDP. Leading the team in massive data ingestion, processing, governance, and downstream distribution for global regulatory compliance operations.')

add_heading('Key Project 1: RCDP Architecture Evolution & Data Governance', level=4)
p = doc.add_paragraph()
p.add_run('Context: ').bold = True
p.add_run('Building the unified cloud data foundation for HSBC\'s R&C function and leading the strategic migration to decommission the legacy on-prem Oracle CDR.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Dual-Track Data Ingestion Architecture: ').bold = True
p.add_run('Designed a dual-track ETL architecture. Implemented a Streaming mode via Pub/Sub + Dataflow for real-time business needs, and a Batch mode via Cloud Scheduler + Dataflow + Landing Buckets for massive historical and periodic data, significantly improving throughput and resilience.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Financial-Grade Data Governance: ').bold = True
p.add_run('Architected an automated Entitlement Engine to meet strict audit requirements. Automatically generated BigQuery Auth Views based on physical table data dictionaries and sensitivity classifications (Public/Internal/Restricted), achieving physical isolation between data storage and data consumption to enforce PoLP.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Diversified Data Provisioning: ').bold = True
p.add_run('Provided low-latency datasets directly via Auth Views for CMLP (Machine Learning Platform) and MI systems. Built highly concurrent, auto-scaling REST API microservices (Cloud Run + Java Spring Boot / Python FastAPI) for general downstream RC applications.')

add_heading('Key Project 2: MLOps Infrastructure & Rapid2 Data Pipeline', level=4)
p = doc.add_paragraph()
p.add_run('Context: ').bold = True
p.add_run('Empowering data science teams with a stable big data processing platform and end-to-end model execution environment.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('End-to-End Streaming ETL: ').bold = True
p.add_run('Constructed complex real-time data flows connecting the Rapid2 business system, RCDP, and the Model Execution Env. Consumed massive compliance alerts via Pub/Sub, cleansed data into the Conformed Layer, triggered real-time model execution, and built a Reverse Flow to instantly write AI decisions back to the upstream system.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cloud Infrastructure & Automated MLOps: ').bold = True
p.add_run('Architected a Kubernetes-based Model Execution Env from scratch. Established automated CI/CD pipelines enabling data scientists to seamlessly test and deploy models trained on CMLP (Kubeflow/Jupyter) into production, completing the last mile of ML engineering.')

add_heading('Previous Management & Technical Experience', level=3)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Team Manager (China) - GCMS (2021.08 - 2023.04): ').bold = True
p.add_run('Led a 9-person Agile team, implementing deep DevOps transformations (Blue-Green deployment, Feature Toggles) and Jenkins/Ansible CI/CD pipelines. Managed Spring Cloud microservices migration to ensure high availability for critical regulatory systems.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Senior Software/System Engineer - AIA & TCS (Client: HSBC) (2011.03 - 2018.06): ').bold = True
p.add_run('Spent 7 years as a core outsourced expert providing foundational system architecture support. Led the architectural upgrade and high-concurrency performance tuning of the core Group Insurance system. Developed HSBC\'s global core trading systems, acquiring deep expertise in underlying financial logic and highly available system design.')


add_heading('Open Source & Personal Projects', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Agentic Web Chat-App (Full-Stack & Cloud-Native): ').bold = True
p.add_run('Cloud-native application deployed on GKE and Cloud Run integrating MCP Servers. Utilized Cloud SQL for persistent data storage and Auth0 for OAuth2 security, demonstrating comprehensive full-stack and cloud deployment (Cloud Build CI/CD) capabilities. (Repo: github.com/nvd11/askc-backend)')

add_heading('Education', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Bachelor of Computer Science').bold = True
p.add_run(', Guangdong University of Foreign Studies | 2003.09 - 2007.07')

doc.save('/home/gateman/.openclaw/workspace/Jason_Pan_Resume_Data_Cloud_V4_EN.docx')
