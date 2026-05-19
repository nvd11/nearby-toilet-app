import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h

# Header
h = add_heading('Jason Pan (Wenlin Pan)', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Target Roles: ').bold = True
p.add_run('Model Validation Director / Head of AI Safety & Governance / Chief AI Architect\n')
p.add_run('Location: ').bold = True
p.add_run('Guangzhou, China | ')
p.add_run('Experience: ').bold = True
p.add_run('17 Years (10+ Years in FinTech Risk & Compliance, Deep Expertise in AI Governance & MLOps)\n')
p.add_run('Education: ').bold = True
p.add_run('Bachelor of Computer Science, Guangdong University of Foreign Studies')

add_heading('Executive Summary', level=2)
p = doc.add_paragraph('Senior FinTech and AI Architecture Expert with over 17 years of software development and architecture experience, including 10+ years specializing in Risk & Compliance.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Enterprise AI Architecture & Governance: ').bold = True
p.add_run('Focused on the highly secure deployment of GenAI and Large Language Models (LLMs) in the financial sector. Led the full-lifecycle governance of an enterprise RAG Agent (ADA) and AI-driven compliance risk system (Rapid2). Expert in non-intrusive AI orchestration, successfully reducing compliance ticket response times from days to under 1 hour.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Model Risk Management (MRM) & Validation: ').bold = True
p.add_run('Deep understanding of global financial regulatory requirements (SR 11-7, Basel, GDPR). Designed and implemented automated model evaluation pipelines and monitoring loops. Leveraged LLM-as-a-Judge to build a multi-dimensional AI safety metrics system covering Completeness, Faithfulness, Toxicity, and Jailbreak Interception.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('MLOps & Cloud-Native Engineering: ').bold = True
p.add_run('Geek-level practical capability in cloud-native engineering, proficient in the GCP ecosystem (BigQuery, Dataflow, Kubeflow, GKE). Built a Model Execution Environment from scratch bridging Data Science and Business systems, automating the end-to-end flow from data cleansing and feature engineering to one-click model deployment and monitoring.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cutting-edge AI Pioneer: ').bold = True
p.add_run('Possesses the rare hybrid advantage of "understanding strategy while mastering low-level code." Proficient in LangChain and ReAct paradigms, continuously contributing to open-source communities with smart agents based on the MCP (Model Context Protocol) and GitOps automated code review tools.')

add_heading('Core Competencies', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('AI Strategy & Governance: ').bold = True
p.add_run('Model Risk Management (MRM), GenAI/LLM Evaluation Metrics, Responsible AI, Jailbreak & Sensitivity Interception, Data Bias Mitigation.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Enterprise AI Architecture: ').bold = True
p.add_run('Non-intrusive AI Orchestration Service, RAG Architecture, Vector Databases (AlloyDB), MCP Server Development & Integration, LangChain.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('MLOps & Cloud-Native: ').bold = True
p.add_run('GCP (Dataflow, BigQuery, Pub/Sub, Cloud Run), Kubernetes (GKE), Kubeflow ML Platform, Jenkins/Terraform CI/CD.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Software & Data Engineering: ').bold = True
p.add_run('Python (Asynchronous High-Performance), Java (Spring Cloud Microservices), Streaming/Batch ETL, High-Concurrency Architecture Design.')


add_heading('Professional Experience', level=2)
add_heading('HSBC Software Development (Guangdong)', level=3)
p = doc.add_paragraph()
p.add_run('Role: ').bold = True
p.add_run('Technical Lead & Engineering Manager - Regulatory Compliance Data Platform (RCDP) | ')
p.add_run('Date: ').bold = True
p.add_run('2018.06 - Present\n')
p.add_run('Core Responsibility: ').bold = True
p.add_run('Fully responsible for the technical strategy and architecture evolution of the RCDP platform. Leading cross-functional teams (Data Science, Risk & Compliance, Engineering) to implement enterprise AI projects and compliance data systems.')

add_heading('Key Project 1: RCDP & Ask Compliance Digital Assistant (ADA)', level=4)
p = doc.add_paragraph()
p.add_run('Context: ').bold = True
p.add_run('Integrating GenAI to empower traditional compliance consulting systems, reducing the burden on compliance experts querying massive cross-border policy documents.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Enterprise AI Architecture & Compliance Interception: ').bold = True
p.add_run('Designed and implemented the core Orchestration Service (AI Orchestration Layer). Before routing user queries to the LLM, applied NLP and BM25 algorithms for intent relevance analysis and data sensitivity interception, ensuring strict adherence to financial info-sec standards and achieving non-intrusive AI empowerment for legacy systems.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Outstanding Business Impact: ').bold = True
p.add_run('Integrated Global Policy & Procedure (GPPS) data into the knowledge base. Upon launch, the AI Assistant drastically reduced the average compliance query response time from "days" to "under 1 hour", massively liberating the productivity of compliance experts.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Model Evaluation & MRM Governance System: ').bold = True
p.add_run('Developed a Model Evaluation API to collect full-trace data (User Query -> AI Response -> User Feedback). Built a multi-dimensional monitoring dashboard based on LLM-as-a-Judge to track in real-time:')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('Core Quality Metrics: ').bold = True
p2.add_run('Completeness, Faithfulness, Accuracy, Cosine Similarity.')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('Safety & Compliance Defense: ').bold = True
p2.add_run('Toxicity detection, PII sensitive data interception, and Jailbreaking Flag defense, successfully establishing a second line of defense for AI outputs.')


add_heading('Key Project 2: Rapid2 Intelligent Compliance Alert Triage (Model1 Project)', level=4)
p = doc.add_paragraph()
p.add_run('Context: ').bold = True
p.add_run('Seamlessly integrating Machine Learning capabilities into traditional RC alert systems to achieve automated preliminary classification and processing of massive alerts.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('End-to-End MLOps Pipeline: ').bold = True
p.add_run('Addressing the pain point of Data Scientists ("strong in algorithms, weak in engineering"), led the development of the Model Execution Environment. Established CI/CD pipelines to automate model deployment and end-to-end testing from Kubeflow to production.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Real-Time Data Streams & Reverse Validation: ').bold = True
p.add_run('Built Streaming ETL via Dataflow and Pub/Sub to consume and process upstream alert data in real-time. Simultaneously established a Recommendation Reverse Flow to write model decisions back to business systems instantly.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Model Decision Feedback Mechanism: ').bold = True
p.add_run('Ran daily batch comparisons between "Model Recommendations" and "Human Final Decisions", generating deviation and accuracy reports via an Evaluation Model to provide quantitative foundations for continuous model iteration.')


add_heading('Historical Key Management Achievement: GCMS Agile Transformation', level=4)
p = doc.add_paragraph()
p.add_run('Role: ').bold = True
p.add_run('China Team Manager | ')
p.add_run('Date: ').bold = True
p.add_run('2021.08 - 2023.04')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Led a 9-person Agile team, driving the R&D model transition from traditional Waterfall to a 2-Week Sprint agile delivery.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Implemented deep DevOps transformations (Blue-Green Deployment, Feature Toggles), increasing automated testing coverage from 0 to 70% and reducing testing costs by 50%.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Directed the migration to a microservices architecture and resolved Log4j critical vulnerabilities across 20+ components within 1.5 months with zero production incidents.')

add_heading('Previous Experience', level=3)
p = doc.add_paragraph('Prior to formally joining HSBC, spent 7 years as a core outsourced expert providing foundational system architecture support for AIA and HSBC.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('AIA Information Technology & Wicresoft (2016.12 - 2018.06): ').bold = True
p.add_run('Senior System Engineer. Led architectural upgrades, code reviews, and high-concurrency performance tuning for the core Group Insurance system.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Tata Consultancy Services (Client: HSBC) (2011.03 - 2016.12): ').bold = True
p.add_run('Senior Software Engineer. Deeply involved in the development of HSBC\'s global core trading systems, accumulating profound experience in low-level financial business logic, transaction idempotency, and high-availability system design.')


add_heading('Cutting-Edge AI Agent Architecture (Open Source & Personal Projects)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Agentic Web Chat-App (Full-Stack AI Application):').bold = True
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('Architecture: ').bold = True
p2.add_run('A cloud-native application deployed on GKE and Cloud Run, integrating MCP (Model Context Protocol) Servers and multi-layer Agent Routing technologies to achieve intelligent tool routing.')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('Highlights: ').bold = True
p2.add_run('Leveraged Cloud SQL for AI Memory enhancement, implemented OAuth2 via Auth0, and supported LLM streaming output and real-time Markdown rendering on the frontend. (Repo: github.com/nvd11/askc-backend)')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Automated AI Code Reviewer (GitOps-based AI Review Tool):').bold = True
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('Mechanism: ').bold = True
p2.add_run('Developed a GitHub Webhook service that automatically awakens Gemini LLM and WebAgents upon PR creation to conduct deep code walkthroughs, injecting improvement suggestions directly into PR Comments.')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('Enterprise Expansion: ').bold = True
p2.add_run('Deeply integrated Jira and Confluence MCP toolchains to ensure AI reviews not only check code standards but align with internal business contexts. Successfully promoted and adopted by cross-departmental teams.')

add_heading('Education', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Bachelor of Computer Science').bold = True
p.add_run(', Guangdong University of Foreign Studies | 2003.09 - 2007.07')

doc.save('/home/gateman/.openclaw/workspace/Jason_Pan_Resume_AI_EN.docx')
