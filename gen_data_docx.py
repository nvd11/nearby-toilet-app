import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

# Header
h = add_heading('潘文林 (Jason Pan)', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('目标职位: ').bold = True
p.add_run('数据总监 (Head of Data) / 首席云数据架构师 (Cloud Data Architect) / 资深数据工程师\n')
p.add_run('工作地点: ').bold = True
p.add_run('广州 | ')
p.add_run('工作经验: ').bold = True
p.add_run('17年 (10+年金融风控数据平台架构与 GCP 云原生经验)\n')
p.add_run('学历: ').bold = True
p.add_run('本科 (广东外语外贸大学 - 计算机科学与技术)')

add_heading('职业总结 (Executive Summary)', level=2)
p = doc.add_paragraph('资深金融数据架构与云基础架构专家，拥有超过 17 年的软件研发经验。精通构建基于 GCP (Google Cloud Platform) 的 PB 级企业数据平台，深谙从传统本地数据中心向云原生架构的整体演进战略。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('企业级数据平台与上云迁移: ').bold = True
p.add_run('作为技术负责人，主导汇丰银行监管合规数据平台 (RCDP) 的架构设计与研发。制定并执行“CDR Demise”战略，成功将核心业务从传统 On-Premises (Oracle) 数据仓库平滑迁移至 GCP 现代化云原生数据架构。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('数据治理 (Data Governance) 与信息安全: ').bold = True
p.add_run('深刻理解金融级数据安全合规 (GDPR, 银行保密法)。独立设计并实现自动化授权引擎 (Entitlement Engine)，基于字段级敏感度 (Public/Restricted等) 自动生成 BigQuery Auth Views，构建了严密的细粒度数据访问控制体系。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('复杂流批一体数据工程 (Streaming/Batch ETL): ').bold = True
p.add_run('精通基于 Apache Beam (Dataflow) 的高吞吐量数据处理。针对不同业务时效要求，构建了兼顾低延迟 (Pub/Sub + Dataflow) 与高吞吐 (Cloud Scheduler + Dataflow + Cloud Storage) 的数据管道架构。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('云基础架构 (Cloud Infra) 与 MLOps: ').bold = True
p.add_run('精通 GCP 底层基础设施，涵盖 VPC 网络规划、IAM 安全策略分配及 GKE 容器化部署。从零搭建数据科学家的模型执行环境 (Model Execution Env)，打通从数据清洗、特征工程到模型一键部署与监控的 MLOps 全链路。')

add_heading('核心能力 (Core Competencies)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Data Engineering: ').bold = True
p.add_run('Dataflow (Apache Beam), Streaming/Batch ETL, Pub/Sub, Cloud Storage, 高并发数据接入。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Data Governance: ').bold = True
p.add_run('数据血缘追踪，数据资产分级 (Data Sensitivity Classification)，BigQuery Auth Views，Entitlement Engine (鉴权引擎设计)。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cloud Infra (GCP): ').bold = True
p.add_run('BigQuery, GKE (Kubernetes), Cloud Run, VPC/IAM/Service Accounts, Cloud Scheduler, Terraform (IaC)。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Software Engineering: ').bold = True
p.add_run('Python (异步/多线程), Java (Spring Boot/Spring Cloud), FastAPI, RESTful API 架构, CI/CD (Jenkins/CloudBuild)。')


add_heading('工作经历 (Professional Experience)', level=2)
add_heading('汇丰软件开发 (广东) 有限公司 (HSBC)', level=3)
p = doc.add_paragraph()
p.add_run('职位: ').bold = True
p.add_run('技术研发经理 / ITSO - 监管合规数据平台 (RCDP) | ')
p.add_run('时间: ').bold = True
p.add_run('2018.06 - 至今\n')
p.add_run('核心职责：').bold = True
p.add_run('全面负责 RCDP 数据平台的技术栈选型、基础架构搭建以及核心框架的开发。带领团队负责全球监管合规业务的海量数据接入、处理、治理及下游分发。')

add_heading('核心项目 1: 监管合规数据平台 (RCDP) 架构演进与数据治理', level=4)
p = doc.add_paragraph()
p.add_run('项目背景：').bold = True
p.add_run('旨在构建汇丰 R&C 部门在云端的统一数据底座，并承担取代现有本地 Oracle 传统数据仓库 (CDR) 的战略任务。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('流批一体数据摄取架构 (Data Ingestion)：').bold = True
p.add_run('为适应复杂的上下游系统，设计了双轨制 ETL 架构。针对实时性要求高的业务，采用 Pub/Sub + Dataflow 的 Streaming 模式；针对海量历史与周期性数据，采用 Cloud Scheduler + Dataflow + Landing Bucket 的 Batch 模式，大幅提升数据吞吐率与系统韧性。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('金融级数据治理与安全控制体系：').bold = True
p.add_run('面对严苛的合规审计，主导设计了自动化权限控制引擎 (Entitlement Engine)。通过数据字典定义物理表字段的敏感度级别 (Public/Internal/Restricted/Highly Restricted)，自动构建面向不同受众的 BigQuery Auth Views。实现了数据存储 (Physical Tables) 与数据消费 (Auth Views) 的彻底物理隔离与最小权限原则 (PoLP)。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('多元化下游数据供应 (Data Provisioning)：').bold = True
p.add_run('为满足不同业务场景，构建了灵活的数据暴露机制。对于 CMLP (机器学习平台) 和 MI (管理信息系统)，通过 Auth Views 直接提供低延迟的数据集；针对一般 RC 业务系统，基于 Cloud Run + Java Spring Boot / Python FastAPI 构建了高并发、自动扩缩容的 REST API 微服务群。')

add_heading('核心项目 2: MLOps 基础架构与 Rapid2 数据管道闭环', level=4)
p = doc.add_paragraph()
p.add_run('项目背景：').bold = True
p.add_run('赋能数据科学家团队，为其提供稳定的大数据处理平台与端到端的模型运行环境。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('端到端实时数据管道 (Streaming ETL)：').bold = True
p.add_run('构建了连接 Rapid2 业务系统、RCDP 数据层与模型执行环境的复杂实时数据流。通过 Pub/Sub 接收海量合规警报，经 Streaming ETL 清洗后落入 Conformed Layer，并实时触发模型计算，最后建立 Reverse Flow 将 AI 决策结果即时回写至上游业务系统。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('云基础架构与自动化 MLOps：').bold = True
p.add_run('从零规划并搭建基于 Kubernetes 的 Model Execution Env。构建自动化 CI/CD 流水线，使得数据科学家能够将其在 CMLP (Kubeflow/Jupyter) 训练的模型，一键完成自动化测试与生产环境部署，彻底打通了算法到工程落地的最后一公里。')

add_heading('历史核心管理成就：全球冲突管理系统 (GCMS) 敏捷转型', level=4)
p = doc.add_paragraph()
p.add_run('角色: ').bold = True
p.add_run('中国区团队经理 (Team Manager) | ')
p.add_run('时间: ').bold = True
p.add_run('2021.08 - 2023.04')
p = doc.add_paragraph(style='List Bullet')
p.add_run('带领 9 人敏捷团队，实施深度 DevOps 改造（蓝绿部署、Feature Toggles），引入 CI/CD (Jenkins/Ansible) 构建自动化部署流水线。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('领导系统微服务化 (Spring Cloud) 迁移，大幅降低系统耦合度，保障了关键监管业务的高可用性。')


add_heading('前沿技术作品 (Open Source / Side Projects)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Agentic Web Chat-App (全栈架构与云原生部署): ').bold = True
p.add_run('基于 GKE 与 Cloud Run 部署的云原生应用，集成 MCP Server。利用 Cloud SQL 数据库实现数据持久化，通过 Auth0 实现 OAuth2 安全鉴权，展示了全栈开发与云端部署（CloudBuild CI/CD）的综合能力。(Repo: github.com/nvd11/askc-backend)')

add_heading('教育背景 (Education)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('广东外语外贸大学').bold = True
p.add_run(' - 计算机科学与技术 (本科) | 2003.09 - 2007.07')

doc.save('/home/gateman/.openclaw/workspace/Jason_Pan_Resume_Data_Cloud_V4.docx')
