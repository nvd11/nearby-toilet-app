import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

# Header
h = add_heading('潘文林 (Jason Pan)', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('目标职位: ').bold = True
p.add_run('模型验证总监 (Model Validation Director) / AI 安全与治理负责人 / AI 首席架构师\n')
p.add_run('工作地点: ').bold = True
p.add_run('广州 | ')
p.add_run('工作经验: ').bold = True
p.add_run('17年 (10+年金融风控与合规经验，深耕 AI 治理与 MLOps)\n')
p.add_run('学历: ').bold = True
p.add_run('本科 (广东外语外贸大学 - 计算机科学与技术)')

add_heading('职业总结 (Executive Summary)', level=2)
p = doc.add_paragraph('资深金融科技与 AI 架构专家，拥有超过 17 年的软件研发与架构经验，其中 10+ 年专注于金融风险与合规 (Risk & Compliance) 领域。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('企业级 AI 架构与治理: ').bold = True
p.add_run('专注于 GenAI 与大模型在金融领域的高安全性落地。主导企业级 RAG Agent (ADA) 与合规风控 AI 系统 (Rapid2) 的全生命周期治理，精通非侵入式 AI 架构设计，成功将合规工单响应时间从天级别压缩至 1 小时内。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('模型风险管理 (MRM) 与验证: ').bold = True
p.add_run('深度理解全球金融监管合规要求 (SR 11-7, Basel, GDPR)。独立设计并落地自动化模型评估流水线与监控闭环，利用 LLM-as-a-Judge 构建涵盖完整度、忠实度、毒性及越狱拦截的多维 AI 安全指标体系。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('MLOps 与云原生工程: ').bold = True
p.add_run('极客级别的云原生实战能力，精通 GCP 生态 (BigQuery, Dataflow, Kubeflow, GKE)。从零构建面向数据科学家与业务系统的模型执行环境 (Model Execution Env)，实现从数据清洗、特征工程到模型一键部署与监控的端到端自动化。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('前沿 AI 技术开拓者: ').bold = True
p.add_run('具备罕见的“懂战略更懂底层代码”的复合优势。熟练运用 LangChain、ReAct 模式，并在开源社区持续贡献基于 MCP (Model Context Protocol) 协议的智能代理与 GitOps 自动化代码审查工具。')

add_heading('核心能力 (Core Competencies)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('AI 战略与治理: ').bold = True
p.add_run('模型风险管理 (MRM), GenAI/LLM 评估指标构建, 负责任的 AI (Responsible AI), 越狱拦截与敏感度审查, 数据偏见清洗。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('企业级 AI 架构: ').bold = True
p.add_run('非侵入式 AI 编排 (Orchestration Service), RAG 架构, 向量数据库 (AlloyDB), MCP Server 开发与集成, LangChain。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('MLOps 与云原生: ').bold = True
p.add_run('GCP (Dataflow, BigQuery, Pub/Sub, Cloud Run), Kubernetes (GKE), Kubeflow 机器学习平台, Jenkins/Terraform CI/CD。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('软件工程与数据架构: ').bold = True
p.add_run('Python (异步高性能开发), Java (Spring Cloud 微服务), 流批一体数据处理 (Streaming/Batch ETL), 高并发架构设计。')


add_heading('工作经历 (Professional Experience)', level=2)
add_heading('汇丰软件开发 (广东) 有限公司 (HSBC)', level=3)
p = doc.add_paragraph()
p.add_run('职位: ').bold = True
p.add_run('技术研发经理 / ITSO - 监管合规数据平台 (RCDP) | ')
p.add_run('时间: ').bold = True
p.add_run('2018.06 - 至今\n')
p.add_run('核心职责：').bold = True
p.add_run('全面负责 RCDP 平台的技术战略选型与架构演进，主导跨部门 (数据科学、风控合规、工程架构) 的企业级 AI 项目落地与合规数据体系建设。')

add_heading('核心项目 1: 监管合规数据平台 (RCDP) 与 Ask Compliance AI 助手 (ADA)', level=4)
p = doc.add_paragraph()
p.add_run('项目背景：').bold = True
p.add_run('引入 GenAI 赋能传统合规咨询系统，减轻合规专家查阅海量跨国政策文档的成本。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('企业级 AI 架构设计与合规拦截：').bold = True
p.add_run('设计并落地了核心的 Orchestration Service (AI 编排层)。在将用户提问路由至大模型前，基于 NLP 与 BM25 算法进行意图相关性分析与数据敏感度拦截，确保符合金融信息安全标准；实现对传统系统的非侵入式 AI 赋能。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('卓越的业务价值：').bold = True
p.add_run('知识库融合了全球政策数据 (GPPS)，AI 助手上线后，成功将合规查询工单的平均响应时间从“天级别”断崖式降至“1小时以内”，极大释放了合规专家的生产力。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('模型评估与 MRM 治理体系：').bold = True
p.add_run('开发 Model Evaluation API，收集“用户提问-AI回答-用户反馈”全链路数据。构建了基于 LLM-as-a-Judge 的多维监控仪表盘，实时追踪：')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('核心质量指标: ').bold = True
p2.add_run('完整性 (Completeness), 忠实度 (Faithfulness), 准确率 (Accuracy), 语义相似度 (Cosine Similarity)。')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('安全与合规防御: ').bold = True
p2.add_run('毒性检测 (Toxicity), PII 敏感数据拦截, 越狱攻击防御 (Jailbreaking Flag)，成功筑起 AI 输出的第二道防线。')


add_heading('核心项目 2: Rapid2 智能合规警报分流系统 (Model1 Project)', level=4)
p = doc.add_paragraph()
p.add_run('项目背景：').bold = True
p.add_run('将机器学习能力无缝集成至传统 RC 警报系统，实现海量警报的自动化初步分类与处理。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('构建端到端 MLOps 闭环：').bold = True
p.add_run('针对数据科学家“懂算法、弱工程”的痛点，主导开发了 Model Execution Environment (模型执行环境)。建立 CI/CD 流水线，实现从 Kubeflow 到生产环境的模型自动化部署与端到端测试。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('实时数据流与反向验证：').bold = True
p.add_run('基于 Dataflow 与 Pub/Sub 构建 Streaming ETL，实时消费并处理上游警报数据；同时建立 Recommendation Reverse Flow 将模型决策实时写回业务系统。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('模型决策自反馈机制：').bold = True
p.add_run('每日批量对比“模型推荐结论”与“合规人员最终决策”，通过 Evaluation Model 生成偏差与准确度报告，为模型的持续迭代提供量化依据。')


add_heading('历史核心管理成就：全球冲突管理系统 (GCMS) 敏捷转型', level=4)
p = doc.add_paragraph()
p.add_run('角色: ').bold = True
p.add_run('中国区团队经理 (Team Manager) | ')
p.add_run('时间: ').bold = True
p.add_run('2021.08 - 2023.04')
p = doc.add_paragraph(style='List Bullet')
p.add_run('带领 9 人敏捷团队，推动研发模式从传统瀑布向 2-Week Sprint 敏捷交付转型。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('实施深度 DevOps 改造（蓝绿部署、Feature Toggles），将自动化测试覆盖率从 0 提升至 70%，测试成本降低 50%。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('领导微服务架构迁移，并在 1.5 个月内无生产事故地完成了 20+ 个组件的 Log4j 高危漏洞修复。')

add_heading('过往履历简述 (Previous Experience)', level=3)
p = doc.add_paragraph('在加入汇丰正式编制前，长达 7 年以核心外派专家身份为友邦保险 (AIA) 与汇丰银行提供底层系统架构支持。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('友邦资讯科技 (2017.11 - 2018.06) / 上海微创 (2016.12 - 2017.11): ').bold = True
p.add_run('担任高级系统工程师，主导核心团险系统 (Group Insurance) 的架构升级、代码审查与高并发性能调优。')
p = doc.add_paragraph(style='List Bullet')
p.add_run('塔塔信息技术 TCS - 汇丰项目 (2011.03 - 2016.12): ').bold = True
p.add_run('担任高级软件工程师，深度参与汇丰全球核心交易系统开发，积累了深厚的底层金融业务逻辑、交易幂等性及高可用系统设计经验。')


add_heading('前沿 AI Agent 架构作品 (Open Source & Personal Projects)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Agentic Web Chat-App (全栈 AI 聊天应用):').bold = True
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('架构: ').bold = True
p2.add_run('基于 GKE 与 Cloud Run 部署的云原生应用，集成 MCP Server (Model Context Protocol) 与多层 Agent Routing 技术，实现智能工具路由。')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('亮点: ').bold = True
p2.add_run('利用 Cloud SQL 实现 AI Memory 增强，采用 Auth0 进行安全鉴权，前端支持 LLM 流式输出与 Markdown 实时渲染。(Repo: github.com/nvd11/askc-backend)')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Automated AI Code Reviewer (基于 GitOps 的 AI 审查工具):').bold = True
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('机制: ').bold = True
p2.add_run('开发 GitHub Webhook 服务，在 PR 创建时自动唤醒 Gemini LLM 及 WebAgent 进行深度代码走查，并将改进建议直接打入 PR Comment。')
p2 = doc.add_paragraph(style='List Bullet 2')
p2.add_run('企业级扩展: ').bold = True
p2.add_run('深度集成 Jira 和 Confluence 的 MCP 工具链，确保 AI 审查不仅关注代码规范，更对齐企业内部的业务上下文。该工具已成功推广至跨部门团队。')

add_heading('教育背景 (Education)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('广东外语外贸大学').bold = True
p.add_run(' - 计算机科学与技术 (本科) | 2003.09 - 2007.07')

doc.save('/home/gateman/.openclaw/workspace/Jason_Pan_Resume_V3_Native.docx')
